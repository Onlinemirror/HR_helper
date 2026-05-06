"""
Все операции с Google Sheets, Google Drive и Google Docs.
- Google Sheets: сервисный аккаунт
- Google Drive (файлы пользователя): OAuth2 токен
- Google Docs (замена переменных): сервисный аккаунт
"""
import io
import mimetypes
import os
import re
import threading
from datetime import datetime

# Блокировки для get_or_create_subfolder — предотвращают race condition
# когда несколько загрузок стартуют одновременно и каждая создаёт свою папку
_subfolder_locks: dict[str, threading.Lock] = {}
_subfolder_locks_meta = threading.Lock()  # защищает сам словарь блокировок

import gspread
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials as OAuthCredentials
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload

import config

# ── pymorphy2 (склонение ФИО) — опциональная зависимость ─────────────────────
try:
    import pymorphy3 as _pymorphy3
    _morph = _pymorphy3.MorphAnalyzer()
    _MORPH_OK = True
except ImportError:
    _MORPH_OK = False

SCOPES_SA = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
]
DRIVE_TOKEN_FILE = "credentials/drive_token.json"


# ── Клиенты ────────────────────────────────────────────────────────────────────

def _sa_credentials() -> Credentials:
    return Credentials.from_service_account_file(
        config.GOOGLE_CREDENTIALS_FILE, scopes=SCOPES_SA
    )


def _oauth_credentials() -> OAuthCredentials:
    if not os.path.exists(DRIVE_TOKEN_FILE):
        raise FileNotFoundError(
            f"Файл '{DRIVE_TOKEN_FILE}' не найден. "
            "Запустите 'python setup_drive_auth.py' для авторизации Google Drive."
        )
    creds = OAuthCredentials.from_authorized_user_file(
        DRIVE_TOKEN_FILE, ["https://www.googleapis.com/auth/drive"]
    )
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with open(DRIVE_TOKEN_FILE, "w") as f:
            f.write(creds.to_json())
    return creds


def _gspread_client() -> gspread.Client:
    return gspread.authorize(_sa_credentials())


def _drive_service():
    """Drive от имени пользователя (OAuth2) — для загрузки файлов."""
    return build("drive", "v3", credentials=_oauth_credentials(), cache_discovery=False)


def _drive_sa():
    """Drive от имени сервисного аккаунта — для поиска шаблонов."""
    return build("drive", "v3", credentials=_sa_credentials(), cache_discovery=False)


def _docs_service():
    return build("docs", "v1", credentials=_sa_credentials(), cache_discovery=False)


def _spreadsheet() -> gspread.Spreadsheet:
    return _gspread_client().open_by_key(config.SPREADSHEET_ID)


def _master_sheet() -> gspread.Worksheet:
    return _spreadsheet().worksheet(config.MASTER_SHEET)


def _log_sheet() -> gspread.Worksheet:
    return _spreadsheet().worksheet(config.LOG_SHEET)


def _get_or_create_sheet(name: str) -> gspread.Worksheet:
    ss = _spreadsheet()
    try:
        return ss.worksheet(name)
    except gspread.WorksheetNotFound:
        return ss.add_worksheet(title=name, rows=500, cols=20)


# ── Вспомогательные ───────────────────────────────────────────────────────────

def extract_folder_id(link_or_id: str) -> str:
    match = re.search(r"folders/([a-zA-Z0-9_-]+)", link_or_id)
    if match:
        return match.group(1)
    return link_or_id.strip()


def _find_folder(drive, name: str, parent_id: str) -> str | None:
    q = (f"name='{name}' and mimeType='application/vnd.google-apps.folder' "
         f"and '{parent_id}' in parents and trashed=false")
    res = drive.files().list(
        q=q, fields="files(id)", supportsAllDrives=True,
        includeItemsFromAllDrives=True
    ).execute()
    files = res.get("files", [])
    return files[0]["id"] if files else None


def _find_or_create_folder_sa(name: str, parent_id: str) -> str:
    """Ищет или создаёт папку через сервисный аккаунт."""
    drive = _drive_sa()
    existing = _find_folder(drive, name, parent_id)
    if existing:
        return existing
    meta = {
        "name": name,
        "mimeType": "application/vnd.google-apps.folder",
        "parents": [parent_id],
    }
    folder = drive.files().create(
        body=meta, fields="id", supportsAllDrives=True
    ).execute()
    return folder["id"]


def _find_file_by_name(drive, name: str, parent_id: str) -> tuple[str, str] | None:
    """
    Ищет файл по имени (с расширением или без).
    Возвращает (file_id, mime_type) или None.
    """
    for candidate in [name, name + ".docx", name + ".doc"]:
        q = f"name='{candidate}' and '{parent_id}' in parents and trashed=false"
        res = drive.files().list(
            q=q, fields="files(id,mimeType)", supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = res.get("files", [])
        if files:
            return files[0]["id"], files[0]["mimeType"]
    return None


def _replace_in_doc(doc_id: str, variables: dict) -> None:
    docs = _docs_service()
    requests = [
        {
            "replaceAllText": {
                "containsText": {"text": key, "matchCase": True},
                "replaceText": str(value) if value else "",
            }
        }
        for key, value in variables.items()
    ]
    if requests:
        docs.documents().batchUpdate(
            documentId=doc_id, body={"requests": requests}
        ).execute()


def _replace_in_docx_bytes(content: bytes, variables: dict) -> bytes:
    """
    Заменяет {{ПЕРЕМЕННЫЕ}} в .docx через python-docx.
    Обрабатывает случай, когда Word разбивает маркер на несколько runs:
    объединяет текст всего абзаца, делает замену, кладёт результат в первый run.
    """
    from docx import Document as _DocxDocument

    doc = _DocxDocument(io.BytesIO(content))

    def _replace_para(para) -> None:
        full = "".join(r.text for r in para.runs)
        replaced = full
        for key, val in variables.items():
            replaced = replaced.replace(key, str(val) if val else "")
        if replaced != full and para.runs:
            para.runs[0].text = replaced
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        _replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_prikaz_type(template_name: str) -> str:
    t = template_name.lower()
    if "прием" in t or "приём" in t: return "Приём"
    if "увольн" in t:  return "Увольнение"
    if "перевод" in t: return "Перевод"
    if "отпуск" in t:  return "Отпуск"
    if "оклад" in t:   return "Изменение оклада"
    return "Прочее"


def _build_doc_name(template_name: str, full_name: str, num_tag: str, date: str) -> str:
    base = template_name.replace("Шаблон_", "").replace("_", " ")
    parts = [p for p in [base, full_name, num_tag, date] if p]
    return " — ".join(parts)


# ── Склонение ФИО ────────────────────────────────────────────────────────────

def _decline_word(word: str, case: str) -> str:
    """Склоняет одно слово в нужный падеж через pymorphy2."""
    if not _MORPH_OK or not word:
        return word
    parsed = _morph.parse(word)
    if not parsed:
        return word
    # Ищем разбор с именной семантикой (Имя / Отчество / Фамилия)
    best = None
    for p in parsed:
        grammemes = set(p.tag.grammemes)
        if grammemes & {"Name", "Patr", "Surn"}:
            best = p
            break
    if best is None:
        best = parsed[0]
    inflected = best.inflect({case})
    if inflected:
        w = inflected.word
        return w[0].upper() + w[1:] if w else word
    return word


def _short_name(last: str, first: str, middle: str) -> str:
    """Амиров А.Ш. — фамилия + инициалы для подписей."""
    parts = [last]
    if first:
        parts.append(first[0].upper() + ".")
    if middle:
        parts.append(middle[0].upper() + ".")
    return " ".join(parts)


def decline_full_name(full_name: str, case: str) -> str:
    """
    Склоняет ФИО целиком.

    Падежи pymorphy2:
      nomn — именительный   (кто?)       Иванов Иван Иванович
      gent — родительный    (кого?)      Иванова Ивана Ивановича
      datv — дательный      (кому?)      Иванову Ивану Ивановичу
      accs — винительный    (кого?)      Иванова Ивана Ивановича
      ablt — творительный   (кем?)       Ивановым Иваном Ивановичем
      loct — предложный     (о ком?)     Иванове Иване Ивановиче
    """
    if not _MORPH_OK or not full_name:
        return full_name
    return " ".join(_decline_word(part, case) for part in full_name.strip().split())


# ── Генерация сотрудника ──────────────────────────────────────────────────────

def generate_employee_id(city: str) -> str:
    prefix = config.CITY_PREFIXES.get(city, city[:3].upper())
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    max_num = 0
    for row in all_rows[1:]:
        cell_id = row[config.COL["ID"] - 1] if row else ""
        if cell_id.startswith(prefix + "-"):
            try:
                num = int(cell_id.split("-", 1)[1])
                max_num = max(max_num, num)
            except (IndexError, ValueError):
                pass
    return f"{prefix}-{max_num + 1:04d}"


# ── Поиск сотрудника ──────────────────────────────────────────────────────────

def find_employee(query: str) -> tuple[int, list[str]] | None:
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    iin_col = config.COL["ИИН"] - 1
    id_col  = config.COL["ID"] - 1
    for idx, row in enumerate(all_rows[1:], start=2):
        padded = row + [""] * (config.COL["Примечание"] - len(row))
        if padded[iin_col] == query or padded[id_col] == query:
            return idx, padded
    return None


def find_employees_by_name(query: str) -> list[tuple[int, list[str]]]:
    """Нечёткий поиск по ФИО: возвращает список (row_idx, row_data)."""
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    fio_col = config.COL["Полное ФИО"] - 1
    q = query.strip().lower()
    results = []
    for idx, row in enumerate(all_rows[1:], start=2):
        padded = row + [""] * (config.COL["Примечание"] - len(row))
        if q in padded[fio_col].lower():
            results.append((idx, padded))
    return results


def get_employee_history(employee_id: str) -> list[list[str]]:
    """Возвращает строки лога изменений для данного сотрудника."""
    log = _log_sheet()
    rows = log.get_all_values()
    # Колонка 3 (0-based: 2) — ID сотрудника
    return [r for r in rows[1:] if len(r) > 2 and r[2] == employee_id]


def get_statistics() -> dict:
    """Статистика по активным сотрудникам."""
    rows = get_all_active_employees()
    by_dept: dict[str, int] = {}
    by_city: dict[str, int] = {}
    by_type: dict[str, int] = {}
    dept_col = config.COL["Отдел"] - 1
    city_col = config.COL["Город"] - 1
    type_col = config.COL["Тип договора"] - 1
    for row in rows:
        dept = row[dept_col] or "—"
        city = row[city_col] or "—"
        ctype = row[type_col] or "—"
        by_dept[dept] = by_dept.get(dept, 0) + 1
        by_city[city] = by_city.get(city, 0) + 1
        by_type[ctype] = by_type.get(ctype, 0) + 1
    return {"total": len(rows), "by_dept": by_dept, "by_city": by_city, "by_type": by_type}


def update_employee_field(row_index: int, row_data: list[str], field: str, new_value: str, author: str) -> None:
    """Обновляет одно поле сотрудника и пишет в лог."""
    sheet = _master_sheet()
    col_num = config.COL[field]
    old_value = row_data[col_num - 1]
    sheet.update_cell(row_index, col_num, new_value)
    log = _log_sheet()
    now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    employee_id = row_data[config.COL["ID"] - 1]
    full_name   = row_data[config.COL["Полное ФИО"] - 1]
    log.append_row([now, config.MASTER_SHEET, employee_id, full_name, field, old_value, new_value, author])


def get_all_active_employees() -> list[list[str]]:
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    status_col = config.COL["Статус"] - 1
    result = []
    for row in all_rows[1:]:
        padded = row + [""] * (config.COL["Примечание"] - len(row))
        if padded[status_col] in ("Активный", "Испытательный срок", "Работает"):
            result.append(padded)
    return result


# ── Добавление сотрудника ─────────────────────────────────────────────────────

def add_employee(data: dict, employee_id: str, drive_folder_link: str) -> None:
    sheet = _master_sheet()
    hire_date = datetime.now().strftime("%d.%m.%Y")
    parts = [data["last_name"], data["first_name"]]
    if data.get("middle_name"):
        parts.append(data["middle_name"])
    full_name = " ".join(parts)
    row = [""] * config.COL["Примечание"]
    row[config.COL["ID"] - 1]                   = employee_id
    row[config.COL["Фамилия"] - 1]              = data["last_name"]
    row[config.COL["Имя"] - 1]                  = data["first_name"]
    row[config.COL["Отчество"] - 1]             = data.get("middle_name", "")
    row[config.COL["Полное ФИО"] - 1]           = full_name
    row[config.COL["Дата рождения"] - 1]        = data.get("birth_date", "")
    row[config.COL["ИИН"] - 1]                  = data["iin"]
    row[config.COL["Город"] - 1]                = data["city"]
    row[config.COL["Отдел"] - 1]                = data["department"]
    row[config.COL["Должность"] - 1]            = data["position"]
    row[config.COL["Тип договора"] - 1]         = data.get("contract_type", "")
    row[config.COL["ЧК/ТОО"] - 1]              = data.get("legal_entity", "")
    row[config.COL["Дата приёма"] - 1]          = hire_date
    row[config.COL["Статус"] - 1]               = "Активный"
    row[config.COL["Телефон"] - 1]              = data.get("phone", "")
    row[config.COL["Email"] - 1]                = data.get("email", "")
    row[config.COL["Руководитель"] - 1]         = data.get("manager", "")
    row[config.COL["Папка Drive (ссылка)"] - 1] = drive_folder_link
    sheet.append_row(row, value_input_option="USER_ENTERED")


# ── Увольнение ────────────────────────────────────────────────────────────────

def fire_employee(row_index: int, row_data: list[str], author: str, fire_date: str | None = None) -> None:
    sheet = _master_sheet()
    # Если дата передана из хендлера — используем её, иначе берём сегодня
    fire_date     = fire_date or datetime.now().strftime("%d.%m.%Y")
    old_status    = row_data[config.COL["Статус"] - 1]
    old_fire_date = row_data[config.COL["Дата увольнения"] - 1]
    employee_id   = row_data[config.COL["ID"] - 1]
    full_name     = row_data[config.COL["Полное ФИО"] - 1]
    sheet.update_cell(row_index, config.COL["Статус"],          "Уволен")
    sheet.update_cell(row_index, config.COL["Дата увольнения"], fire_date)
    log = _log_sheet()
    now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    log.append_row([now, config.MASTER_SHEET, employee_id, full_name,
                    "Статус", old_status, "Уволен", author])
    log.append_row([now, config.MASTER_SHEET, employee_id, full_name,
                    "Дата увольнения", old_fire_date, fire_date, author])


# ── Drive (загрузка файлов пользователя) ─────────────────────────────────────

def create_drive_folder(folder_name: str, parent_id: str) -> tuple[str, str]:
    service = _drive_service()
    metadata = {
        "name": folder_name,
        "mimeType": "application/vnd.google-apps.folder",
        "parents": [parent_id],
    }
    folder = service.files().create(
        body=metadata, fields="id,webViewLink", supportsAllDrives=True
    ).execute()
    return folder["id"], folder["webViewLink"]


def sync_drive_folders() -> dict:
    """
    Проходит по всем строкам мастер-листа.
    Если у сотрудника нет папки на Drive — создаёт и записывает ссылку.
    Возвращает {"created": [str, ...], "skipped": int, "errors": [str, ...]}
    """
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()

    created = []
    skipped = 0
    errors  = []

    for idx, row in enumerate(all_rows[1:], start=2):   # start=2: пропускаем заголовок
        padded     = row + [""] * (config.COL["Примечание"] - len(row))
        emp_id     = padded[config.COL["ID"] - 1].strip()
        full_name  = padded[config.COL["Полное ФИО"] - 1].strip()
        drive_link = padded[config.COL["Папка Drive (ссылка)"] - 1].strip()

        # Пропускаем пустые строки
        if not emp_id or not full_name:
            continue

        # Папка уже есть — пропускаем
        if drive_link:
            skipped += 1
            continue

        try:
            folder_name = f"{emp_id} {full_name}"
            _, folder_link = create_drive_folder(folder_name, config.HR_DRIVE_FOLDER_ID)
            sheet.update_cell(idx, config.COL["Папка Drive (ссылка)"], folder_link)
            created.append(f"{emp_id} — {full_name}")
        except Exception as e:
            errors.append(f"{emp_id} — {full_name}: {e}")

    return {"created": created, "skipped": skipped, "errors": errors}


def list_employee_subfolders(folder_id: str) -> list[dict]:
    """
    Возвращает список подпапок в папке сотрудника.
    Каждый элемент: {"id": ..., "name": ...}
    Первым добавляется псевдо-элемент «Корневая папка» (сам folder_id).
    """
    service = _drive_service()
    q = (f"'{folder_id}' in parents "
         f"and mimeType='application/vnd.google-apps.folder' "
         f"and trashed=false")
    res = service.files().list(
        q=q, fields="files(id,name)", orderBy="name",
        supportsAllDrives=True, includeItemsFromAllDrives=True,
    ).execute()
    subfolders = res.get("files", [])
    # Корневая папка идёт первой, чтобы можно было забрать файлы не в подпапке
    return [{"id": folder_id, "name": "📂 Корневая папка"}] + subfolders


def list_files_in_folder(folder_id: str) -> list[dict]:
    """
    Возвращает список файлов (не папок) в указанной папке.
    Каждый элемент: {"id": ..., "name": ..., "mimeType": ..., "size": int}
    """
    service = _drive_service()
    q = (f"'{folder_id}' in parents "
         f"and mimeType!='application/vnd.google-apps.folder' "
         f"and trashed=false")
    res = service.files().list(
        q=q,
        fields="files(id,name,mimeType,size,webViewLink)",
        orderBy="name",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute()
    files = res.get("files", [])
    # size может отсутствовать у Google-native документов
    for f in files:
        f.setdefault("size", 0)
    return files


def download_file_bytes(file_id: str, mime_type: str) -> bytes:
    """
    Скачивает файл и возвращает bytes.
    Google Docs / Sheets / Slides экспортируются в PDF.
    """
    service = _drive_service()
    GOOGLE_NATIVE = {
        "application/vnd.google-apps.document",
        "application/vnd.google-apps.spreadsheet",
        "application/vnd.google-apps.presentation",
    }
    if mime_type in GOOGLE_NATIVE:
        return service.files().export(
            fileId=file_id, mimeType="application/pdf"
        ).execute()
    return service.files().get_media(fileId=file_id).execute()


def get_or_create_subfolder(parent_folder_id: str, subfolder_name: str) -> str:
    """
    Находит или создаёт подпапку категории внутри папки сотрудника.
    Блокировка на (parent_folder_id, subfolder_name) предотвращает race condition:
    несколько одновременных загрузок не создадут дубли папки.
    """
    lock_key = f"{parent_folder_id}:{subfolder_name}"
    with _subfolder_locks_meta:
        if lock_key not in _subfolder_locks:
            _subfolder_locks[lock_key] = threading.Lock()
    lock = _subfolder_locks[lock_key]

    with lock:
        drive = _drive_service()
        # Получаем все подпапки родительской папки одним запросом и ищем в Python
        res = drive.files().list(
            q=(f"'{parent_folder_id}' in parents"
               f" and mimeType='application/vnd.google-apps.folder'"
               f" and trashed=false"),
            fields="files(id,name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
            pageSize=100,
        ).execute()
        for f in res.get("files", []):
            if f["name"] == subfolder_name:
                return f["id"]
        # Не нашли — создаём
        meta = {
            "name": subfolder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [parent_folder_id],
        }
        folder = drive.files().create(
            body=meta, fields="id", supportsAllDrives=True
        ).execute()
        return folder["id"]


def upload_file_to_drive(local_path: str, file_name: str, folder_id: str) -> str:
    service = _drive_service()
    mime_type, _ = mimetypes.guess_type(local_path)
    mime_type = mime_type or "application/octet-stream"
    metadata = {"name": file_name, "parents": [folder_id]}
    media = MediaFileUpload(local_path, mimetype=mime_type, resumable=True)
    uploaded = service.files().create(
        body=metadata, media_body=media,
        fields="id,webViewLink", supportsAllDrives=True
    ).execute()
    return uploaded["webViewLink"]


# ── Нумерация документов ──────────────────────────────────────────────────────

def _max_serial(values: list[str], year: int) -> int:
    """
    Из списка строк вида «№ 007-ЛС/2026-ЧК», «ТД-ТОО-2026-012», «СПР-2026-003»
    извлекает максимальный порядковый номер для заданного года.
    Работает надёжнее подсчёта строк: дубли и удалённые строки не сбивают счётчик.
    """
    import re
    year_str = str(year)
    max_num = 0
    for val in values:
        if year_str not in val:
            continue
        # ищем все последовательности цифр и берём последнюю (обычно это серийный номер)
        nums = re.findall(r"\d+", val)
        for n in nums:
            if n == year_str:
                continue
            try:
                max_num = max(max_num, int(n))
            except ValueError:
                pass
    return max_num


def get_next_prikaz_number(entity: str) -> str:
    """№ 001-ЛС/2026-ЧК — берёт max номера из колонки, не считает строки."""
    sheet_name = config.SHEET_PRIKAZ_CHK if entity == "ЧК" else config.SHEET_PRIKAZ_TOO
    sheet = _get_or_create_sheet(sheet_name)
    year = datetime.now().year
    all_rows = sheet.get_all_values()
    # Приказы_ТОО: номер в col 3 (idx 2); Приказы_ЧК: номер в col 4 (idx 3)
    num_col = 3 if entity == "ЧК" else 2
    nums = [row[num_col] for row in all_rows[1:] if len(row) > num_col and row[num_col]]
    max_num = _max_serial(nums, year)
    return f"№ {max_num + 1:03d}-ЛС/{year}-{entity}"


def get_next_dogovor_number(entity: str, doc_type: str) -> str:
    """ТД-ТОО-2026-001 — берёт max из колонки номеров договоров."""
    sheet_name = config.SHEET_CHK if entity == "ЧК" else config.SHEET_TOO
    sheet = _get_or_create_sheet(sheet_name)
    year = datetime.now().year
    all_rows = sheet.get_all_values()
    # Ищем номера, содержащие doc_type и entity (формат ТД-ТОО-2026-001)
    nums = []
    for row in all_rows[1:]:
        if not row:
            continue
        for cell in row:
            if doc_type in cell and entity in cell and str(year) in cell:
                nums.append(cell)
                break
    max_num = _max_serial(nums, year)
    return f"{doc_type}-{entity}-{year}-{max_num + 1:03d}"


def get_next_doc_number(prefix: str, type_filter: str | None) -> str:
    """СПР-2026-001 / ДОВ-2026-001 / ИСХ-2026-001 — берёт max из колонки."""
    sheet = _get_or_create_sheet(config.SHEET_SPRAVKA)
    year = datetime.now().year
    all_rows = sheet.get_all_values()
    nums = []
    for row in all_rows[1:]:
        if not row:
            continue
        # ищем ячейку с нашим префиксом и годом
        for cell in row:
            if prefix in cell and str(year) in cell:
                nums.append(cell)
                break
    max_num = _max_serial(nums, year)
    return f"{prefix}-{year}-{max_num + 1:03d}"


def lookup_last_dogovor(full_name: str, entity: str) -> str:
    sheet_name = config.SHEET_CHK if entity == "ЧК" else config.SHEET_TOO
    sheet = _get_or_create_sheet(sheet_name)
    all_rows = sheet.get_all_values()
    last = ""
    for row in all_rows[2:]:
        if len(row) > 3 and str(row[3]).strip() == full_name.strip() and row[1]:
            last = str(row[1])
    return last


# ── Генерация документа ───────────────────────────────────────────────────────

def generate_document(
    template_name: str,
    entity: str,
    employee_row: list[str],
    extra_vars: dict | None = None,
    prikaz_num: str | None = None,   # если передан — используется вместо авто-счётчика
) -> dict:
    """
    Главная функция генерации документа.
    Возвращает {"pdf_url": ..., "doc_name": ..., "numbers": {...}}
    """
    # ── Данные сотрудника ──────────────────────────────────────────────────
    id_       = employee_row[config.COL["ID"] - 1]
    last_name = employee_row[config.COL["Фамилия"] - 1]
    first_name= employee_row[config.COL["Имя"] - 1]
    middle    = employee_row[config.COL["Отчество"] - 1]
    full_name = employee_row[config.COL["Полное ФИО"] - 1]
    if not full_name:
        full_name = " ".join(filter(None, [last_name, first_name, middle]))
    position    = employee_row[config.COL["Должность"] - 1]
    department  = employee_row[config.COL["Отдел"] - 1]
    city        = employee_row[config.COL["Город"] - 1]
    iin         = employee_row[config.COL["ИИН"] - 1]
    hire_date   = str(employee_row[config.COL["Дата приёма"] - 1])
    folder_link = employee_row[config.COL["Папка Drive (ссылка)"] - 1]

    # ── Тип документа ─────────────────────────────────────────────────────
    t           = template_name.lower()
    is_prikaz   = "приказ" in t
    is_dogovor  = "договор" in t
    is_td       = "_тд" in t
    is_gph      = "_гпх" in t
    is_spravka  = "справк" in t
    is_dov      = "доверен" in t
    is_uved     = "уведом" in t
    tip_dog     = "ТД" if is_td else ("ГПХ" if is_gph else "ТД")

    # ── Номера документов ─────────────────────────────────────────────────
    num_prikaz  = (prikaz_num if prikaz_num
                   else get_next_prikaz_number(entity) if is_prikaz
                   else "")
    num_dogovor = (get_next_dogovor_number(entity, tip_dog) if is_dogovor
                   else lookup_last_dogovor(full_name, entity))
    num_spravka = get_next_doc_number("СПР", "Справка с места работы") if is_spravka else ""
    num_dov     = get_next_doc_number("ДОВ", "Доверенность") if is_dov else ""
    num_ish     = get_next_doc_number("ИСХ", None) if (is_spravka or is_dov or is_uved) else ""

    # ── Находим шаблон (через сервисный аккаунт) ─────────────────────────
    drive_sa = _drive_sa()
    templ_folder_id = _find_or_create_folder_sa(config.TEMPLATES_FOLDER_NAME, config.HR_DRIVE_FOLDER_ID)
    found = _find_file_by_name(drive_sa, template_name, templ_folder_id)
    if not found:
        raise FileNotFoundError(
            f"Шаблон «{template_name}» не найден в папке «{config.TEMPLATES_FOLDER_NAME}»"
        )
    template_file_id, template_mime = found

    # ── Папка сотрудника и подпапка ───────────────────────────────────────
    emp_folder_id = extract_folder_id(folder_link) if folder_link else config.HR_DRIVE_FOLDER_ID
    if is_prikaz:
        sub = "Приказы"
    elif is_dogovor:
        sub = "Договоры"
    elif is_spravka or is_dov or is_uved:
        sub = "Справки"
    else:
        sub = "Прочее"
    dest_folder_id = _find_or_create_folder_sa(sub, emp_folder_id)

    # ── Копируем/конвертируем шаблон в Google Doc ─────────────────────────
    today    = datetime.now().strftime("%d.%m.%Y")
    num_tag  = num_prikaz or (num_dogovor if is_dogovor else "") or num_spravka or num_dov or num_ish
    doc_name = _build_doc_name(template_name, full_name, num_tag, today)

    GDOC_MIME = "application/vnd.google-apps.document"

    # ── Собираем переменные для замены ───────────────────────────────────────
    fire_date_str = employee_row[config.COL["Дата увольнения"] - 1]

    variables = {
        # Базовые данные сотрудника
        "{{ФИО}}":                full_name,
        "{{ДОЛЖНОСТЬ}}":          position,
        "{{ОТДЕЛ}}":              department,
        "{{ГОРОД}}":              city,
        "{{ДАТА_ПРИЁМА}}":        hire_date,
        "{{ДАТА}}":               today,
        "{{ID}}":                 id_,
        "{{ИИН}}":                iin,
        # ФИО в падежах (pymorphy3)
        "{{ФИО_РОД}}":            decline_full_name(full_name, "gent"),
        "{{ФИО_ДАТ}}":            decline_full_name(full_name, "datv"),
        "{{ФИО_ВИН}}":            decline_full_name(full_name, "accs"),
        "{{ФИО_ТВ}}":             decline_full_name(full_name, "ablt"),
        "{{ФИО_ПР}}":             decline_full_name(full_name, "loct"),
        # Компания / документы
        "{{КОМПАНИЯ}}":           config.COMPANY_NAME,
        "{{БИН}}":                config.COMPANY_BIN,
        "{{ГОД}}":                str(datetime.now().year),
        "{{ЮРЛИЦО}}":             entity,
        "{{ТИП_ДОГОВОРА}}":       tip_dog,
        "{{НОМЕР_ПРИКАЗА}}":      num_prikaz,
        "{{НОМЕР_ДОГОВОРА}}":     num_dogovor,
        "{{НОМЕР_СПРАВКИ}}":      num_spravka,
        "{{НОМЕР_ДОВЕРЕННОСТИ}}": num_dov,
        "{{НОМЕР_ИСХ}}":          num_ish,
        "{{ДАТА_УВОЛЬНЕНИЯ}}":    fire_date_str,
        # Дополнительные поля из мастер-листа
        "{{АДРЕС}}":              employee_row[config.COL["Адрес"] - 1],
        "{{РУКОВОДИТЕЛЬ}}":       employee_row[config.COL["Руководитель"] - 1],
        "{{ТЕЛЕФОН}}":            employee_row[config.COL["Телефон"] - 1],
        # Фамилия И.О. — для подписей
        "{{ФИО_КРАТКО}}":         _short_name(last_name, first_name, middle),
    }
    if extra_vars:
        variables.update(extra_vars)

    if template_mime == GDOC_MIME:
        # Нативный Google Doc — копируем, затем заменяем через Docs API
        copied = drive_sa.files().copy(
            fileId=template_file_id,
            body={"name": doc_name, "parents": [dest_folder_id]},
            supportsAllDrives=True,
            fields="id",
        ).execute()
        doc_id = copied["id"]
        # Даём сервисному аккаунту доступ к файлу для Docs API
        sa_email = _sa_credentials().service_account_email
        _drive_service().permissions().create(
            fileId=doc_id,
            body={"type": "user", "role": "writer", "emailAddress": sa_email},
            supportsAllDrives=True,
        ).execute()
        _replace_in_doc(doc_id, variables)
    else:
        # .docx — скачиваем через SA, заменяем локально (python-docx),
        # импортируем готовый .docx через OAuth (квота пользователя)
        content = drive_sa.files().get_media(fileId=template_file_id).execute()
        modified = _replace_in_docx_bytes(content, variables)
        word_mime = ("application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document")
        drive_oauth = _drive_service()
        media = MediaIoBaseUpload(io.BytesIO(modified), mimetype=word_mime, resumable=False)
        imported = drive_oauth.files().create(
            body={"name": doc_name, "mimeType": GDOC_MIME, "parents": [dest_folder_id]},
            media_body=media,
            supportsAllDrives=True,
            fields="id",
        ).execute()
        doc_id = imported["id"]

    # ── Экспортируем в PDF и загружаем через OAuth ─────────────────────────
    drive_oauth = _drive_service()
    pdf_bytes = drive_oauth.files().export(
        fileId=doc_id, mimeType="application/pdf"
    ).execute()

    drive_oauth = _drive_service()
    pdf_media = MediaIoBaseUpload(
        io.BytesIO(pdf_bytes), mimetype="application/pdf", resumable=False
    )
    pdf_file = drive_oauth.files().create(
        body={"name": doc_name + ".pdf", "parents": [dest_folder_id]},
        media_body=pdf_media,
        supportsAllDrives=True,
        fields="id,webViewLink",
    ).execute()
    pdf_url = pdf_file["webViewLink"]

    # Удаляем черновик Google Doc (через OAuth — файл принадлежит пользователю)
    drive_oauth.files().delete(fileId=doc_id, supportsAllDrives=True).execute()

    # ── Пишем в реестр ────────────────────────────────────────────────────
    if is_prikaz:
        _write_prikaz(entity, {
            "num": num_prikaz, "city": city, "type": _get_prikaz_type(template_name),
            "emp_name": full_name, "emp_id": id_, "pos": position,
            "dogovor": num_dogovor, "file_url": pdf_url,
        })
    elif is_dogovor:
        _write_dogovor(entity, {
            "type": tip_dog, "num": num_dogovor, "name": full_name,
            "iin": iin, "subject": position, "dept": department, "file_url": pdf_url,
        })
    elif is_spravka or is_dov or is_uved:
        doc_type = ("Справка с места работы" if is_spravka
                    else "Доверенность" if is_dov else "Уведомление")
        _write_spravka({
            "num": num_ish or num_spravka or num_dov,
            "type": doc_type, "name": full_name,
            "emp_id": id_, "file_url": pdf_url,
        })

    # ── Лог ───────────────────────────────────────────────────────────────
    log = _log_sheet()
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    log.append_row([now_str, config.MASTER_SHEET, id_, full_name,
                    "Документ создан", "", doc_name, "Telegram Bot"])

    return {
        "pdf_url":  pdf_url,
        "doc_name": doc_name,
        "numbers": {
            "prikaz":  num_prikaz,
            "dogovor": num_dogovor if is_dogovor else "",
            "spravka": num_spravka,
            "dov":     num_dov,
            "ish":     num_ish,
        },
    }


# ── Запись в реестры ──────────────────────────────────────────────────────────

def _next_empty_row(sheet: gspread.Worksheet, header_rows: int = 2) -> tuple[int, int]:
    """
    Возвращает (row_1based, serial_number) — номер первой пустой строки
    после заголовков и порядковый номер новой записи.
    """
    all_rows = sheet.get_all_values()
    data_rows = [r for r in all_rows[header_rows:] if any(r)]
    serial = len(data_rows) + 1
    next_row = len(all_rows) + 1  # строка после последней непустой
    # Иногда get_all_values обрезает пустые строки — берём максимум
    next_row = max(next_row, header_rows + len(data_rows) + 1)
    return next_row, serial


def _write_prikaz(entity: str, p: dict) -> None:
    """
    Приказы_ТОО: дата | ФИО | номер приказа | тип | Город | № ТД
    Приказы_ЧК:  дата | ФИО | тип           | номер приказа | Город | Примечание
    """
    sheet = _get_or_create_sheet(
        config.SHEET_PRIKAZ_CHK if entity == "ЧК" else config.SHEET_PRIKAZ_TOO
    )
    today = datetime.now().strftime("%d.%m.%Y")
    row_idx, _ = _next_empty_row(sheet, header_rows=1)
    if entity == "ЧК":
        row = [today, p["emp_name"], p["type"], p["num"], p["city"], ""]
    else:
        row = [today, p["emp_name"], p["num"], p["type"], p["city"], p.get("dogovor", "")]
    sheet.update(
        range_name=f"A{row_idx}",
        values=[row],
        value_input_option="USER_ENTERED",
    )


def _write_dogovor(entity: str, d: dict) -> None:
    sheet = _get_or_create_sheet(
        config.SHEET_CHK if entity == "ЧК" else config.SHEET_TOO
    )
    today = datetime.now().strftime("%d.%m.%Y")
    row_idx, serial = _next_empty_row(sheet, header_rows=2)
    sheet.update(
        range_name=f"A{row_idx}",
        values=[[
            serial, d["num"], d["type"], d["name"], d["iin"],
            d["subject"], d["dept"], "", today, "",
            "", "", "Работает", "", d.get("file_url", ""), ""
        ]],
        value_input_option="USER_ENTERED",
    )


def _write_spravka(s: dict) -> None:
    sheet = _get_or_create_sheet(config.SHEET_SPRAVKA)
    today = datetime.now().strftime("%d.%m.%Y")
    row_idx, _ = _next_empty_row(sheet, header_rows=2)
    sheet.update(
        range_name=f"A{row_idx}",
        values=[[
            s["num"], s["type"], s["name"], s["emp_id"],
            today, "", "", "", s.get("file_url", "")
        ]],
        value_input_option="USER_ENTERED",
    )


# ── Ежедневные проверки ───────────────────────────────────────────────────────

def check_birthdays() -> list[str]:
    """Возвращает список строк-уведомлений о ближайших ДР из мастер-листа."""
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    alerts = []
    fio_col    = config.COL["Полное ФИО"] - 1
    city_col   = config.COL["Город"] - 1
    bd_col     = config.COL["Дата рождения"] - 1
    status_col = config.COL["Статус"] - 1
    for row in all_rows[1:]:
        padded = row + [""] * (config.COL["Примечание"] - len(row))
        if padded[status_col] == "Уволен" or not padded[bd_col]:
            continue
        try:
            bd = datetime.strptime(str(padded[bd_col])[:10], "%d.%m.%Y")
        except ValueError:
            continue
        next_bd = bd.replace(year=today.year)
        if next_bd < today:
            next_bd = next_bd.replace(year=today.year + 1)
        diff = (next_bd - today).days
        if 0 <= diff <= config.DAYS_BEFORE_BDAY:
            name = padded[fio_col]
            city = padded[city_col]
            age  = today.year - bd.year + (0 if next_bd.year > today.year else 0)
            suffix = "🎂 СЕГОДНЯ!" if diff == 0 else f"через {diff} дн."
            alerts.append(f"• {name} ({city}) — {suffix}, исполняется {age} лет")
    return alerts


def check_contract_expiry() -> list[str]:
    """Возвращает список строк об истекающих договорах."""
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    alerts = []
    for sheet_name, entity in [(config.SHEET_TOO, "ТОО"), (config.SHEET_CHK, "ЧК")]:
        try:
            sheet = _get_or_create_sheet(sheet_name)
        except Exception:
            continue
        all_rows = sheet.get_all_values()
        for row in all_rows[2:]:
            if len(row) < 13 or not row[9]:
                continue
            num    = row[1]
            typ    = row[2]
            name   = row[3]
            status = row[12] if len(row) > 12 else ""
            if status not in ("Работает", "Активен", ""):
                continue
            try:
                end_date = datetime.strptime(str(row[9])[:10], "%d.%m.%Y")
            except ValueError:
                continue
            diff = (end_date - today).days
            if 0 <= diff <= config.DAYS_BEFORE_EXPIRY:
                alerts.append(
                    f"• [{entity} {typ}] {name} — {num} | до: "
                    f"{end_date.strftime('%d.%m.%Y')} (через {diff} дн.)"
                )
    return alerts


def check_probation() -> list[str]:
    """Сотрудники, у которых заканчивается испытательный срок (3 мес.)."""
    sheet = _master_sheet()
    all_rows = sheet.get_all_values()
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    alerts = []
    for row in all_rows[1:]:
        padded = row + [""] * (config.COL["Примечание"] - len(row))
        status = padded[config.COL["Статус"] - 1]
        hire   = padded[config.COL["Дата приёма"] - 1]
        if status != "Испытательный срок" or not hire:
            continue
        try:
            hire_dt = datetime.strptime(str(hire)[:10], "%d.%m.%Y")
        except ValueError:
            continue
        end_probation = hire_dt.replace(month=hire_dt.month + 3) if hire_dt.month <= 9 else \
            hire_dt.replace(year=hire_dt.year + 1, month=hire_dt.month - 9)
        diff = (end_probation - today).days
        if 0 <= diff <= 7:
            name = padded[config.COL["Полное ФИО"] - 1]
            pos  = padded[config.COL["Должность"] - 1]
            emp_id = padded[config.COL["ID"] - 1]
            alerts.append(
                f"• {name} ({emp_id}) — {pos} | конец ИС: "
                f"{end_probation.strftime('%d.%m.%Y')} (через {diff} дн.)"
            )
    return alerts


# ── Оценка 360 ────────────────────────────────────────────────────────────────

def save_360_evaluation(
    evaluator_name: str,
    employee_row: list[str],
    scores: dict[str, int],
    strengths: str,
    improve: str,
) -> None:
    sheet = _get_or_create_sheet(config.SHEET_360)
    period = f"Q{(datetime.now().month - 1) // 3 + 1}-{datetime.now().year}"
    emp_name = employee_row[config.COL["Полное ФИО"] - 1]
    emp_id   = employee_row[config.COL["ID"] - 1]
    avg      = round(sum(scores.values()) / len(scores), 2) if scores else 0
    now_str  = datetime.now().strftime("%d.%m.%Y %H:%M")
    sheet.append_row([
        now_str, emp_name, emp_id, period, evaluator_name,
        avg,
        scores.get("quality", ""),
        scores.get("teamwork", ""),
        scores.get("initiative", ""),
        scores.get("communication", ""),
        scores.get("knowledge", ""),
        strengths,
        improve,
    ])
